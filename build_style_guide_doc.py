from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = "Story Nook CSS Editing Guide.docx"


def shade(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def border(paragraph, color="D9B778"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), color)
        p_bdr.append(node)
    p_pr.append(p_bdr)


def add_code(doc, text):
    p = doc.add_paragraph(style="CodeBlock")
    p.add_run(text)
    shade(p, "F7F0DE")
    border(p, "D8BE86")
    return p


def add_find(doc, text):
    p = doc.add_paragraph()
    r = p.add_run("Copy/paste search in style.css: ")
    r.bold = True
    p.add_run(text).font.name = "Consolas"
    return p


def add_tip(doc, text):
    p = doc.add_paragraph(style="Tip")
    p.add_run("Tip: ").bold = True
    p.add_run(text)
    shade(p, "EFE5CE")
    return p


def section(doc, title, intro=None):
    doc.add_heading(title, level=1)
    if intro:
        doc.add_paragraph(intro)


def sub(doc, title, find=None, code=None, notes=None):
    doc.add_heading(title, level=2)
    if find:
        add_find(doc, find)
    if code:
        add_code(doc, code)
    if notes:
        for note in notes:
            doc.add_paragraph(note, style="List Bullet")


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.65)
sec.bottom_margin = Inches(0.65)
sec.left_margin = Inches(0.7)
sec.right_margin = Inches(0.7)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.2)
styles["Normal"].paragraph_format.space_after = Pt(4)

for style_name in ("Heading 1", "Heading 2", "Heading 3"):
    styles[style_name].font.name = "Aptos Display"
styles["Heading 1"].font.size = Pt(18)
styles["Heading 1"].font.color.rgb = RGBColor(74, 45, 24)
styles["Heading 2"].font.size = Pt(13.5)
styles["Heading 2"].font.color.rgb = RGBColor(112, 76, 40)

code_style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
code_style.font.name = "Consolas"
code_style.font.size = Pt(8.4)
code_style.paragraph_format.left_indent = Inches(0.1)
code_style.paragraph_format.right_indent = Inches(0.1)
code_style.paragraph_format.space_before = Pt(3)
code_style.paragraph_format.space_after = Pt(7)

tip_style = styles.add_style("Tip", WD_STYLE_TYPE.PARAGRAPH)
tip_style.font.name = "Aptos"
tip_style.font.size = Pt(9.4)
tip_style.paragraph_format.left_indent = Inches(0.1)
tip_style.paragraph_format.right_indent = Inches(0.1)
tip_style.paragraph_format.space_before = Pt(3)
tip_style.paragraph_format.space_after = Pt(7)

title = doc.add_paragraph()
title_run = title.add_run("The Story Nook CSS Editing Guide")
title_run.bold = True
title_run.font.size = Pt(24)
title_run.font.color.rgb = RGBColor(74, 45, 24)
doc.add_paragraph("A practical map for changing the visible parts of the current Story Nook homepage.")
add_tip(doc, "Most of the newest controls are near the bottom of style.css. When changing something, search for the exact selector or comment shown in this guide, then edit the final matching block.")
add_tip(doc, "If a change does not show up, there is probably a later rule overriding it. Add your change near the bottom of style.css or edit the final override section.")

section(doc, "1. Header", "The header is the top navigation bar containing the main logo, logo 2 wordmark image, quote, login button, and dormant ribbon/menu code.")
sub(
    doc,
    "Header Height And Lower Edge",
    "Header lower edge adjustment",
    """.site-nav,
#mainNav.site-nav {
    padding-bottom: 25px !important;
}

.site-nav.scrolled,
#mainNav.site-nav.scrolled {
    padding-bottom: 12px !important;
}""",
    [
        "Increase padding-bottom to make the header taller at the bottom.",
        "Decrease padding-bottom to make the header shorter.",
        "The scrolled version exists because older code still references it; keep it similar if you want no visual jump."
    ],
)
sub(
    doc,
    "Header Background Gradient",
    "Continuous main-page gradient",
    """:root {
    --nook-page-gradient:
        linear-gradient(180deg,
            #120c08 0%,
            #1b110b 18%,
            #26170e 36%,
            #352113 55%,
            #51341d 76%,
            #74512f 100%);
}""",
    [
        "These colors control the continuous brown gradient used by the header, boxes, and footer.",
        "Change the top colors for the header feel; change lower colors for boxes farther down the page."
    ],
)
sub(
    doc,
    "Main Header Logo",
    "#navLogo .brand-logo-img",
    """#navLogo .brand-logo-img,
#navLogo.brand-lockup.nav-logo .brand-logo-img {
    width: 120px !important;
    height: 120px !important;
    object-fit: cover !important;
    object-position: center center !important;
    transform: translate(15px, 25px) !important;
}""",
    [
        "width and height change the logo size.",
        "transform: translate(X, Y) moves it. Positive X moves right; positive Y moves down.",
        "This uses story-nook-main-logo-new-square.png in index.html."
    ],
)
sub(
    doc,
    "Logo 2 Wordmark Image",
    "Nav wordmark image: use story-nook-logo2.png",
    """#navLogo.brand-lockup.nav-logo .brand-title-img {
    display: block !important;
    position: fixed !important;
    left: 97px !important;
    top: -25px !important;
    height: 196px !important;
}""",
    [
        "left moves logo 2 horizontally.",
        "top moves logo 2 vertically; negative values move it upward.",
        "height controls logo 2 size.",
        "The filename must match index.html exactly: story-nook-logo2.png."
    ],
)
sub(
    doc,
    "Mobile Logo Positions",
    "@media (max-width: 760px)",
    """@media (max-width: 760px) {
    #navLogo .brand-logo-img,
    #navLogo.brand-lockup.nav-logo .brand-logo-img {
        width: 120px !important;
        height: 120px !important;
        transform: translate(14px, 19px) !important;
    }

    #navLogo.brand-lockup.nav-logo .brand-title-img {
        left: 95px !important;
        top: -20px !important;
        height: 180px !important;
    }
}""",
    [
        "This controls phones and narrow screens.",
        "There is also a max-width: 390px block for very small phones."
    ],
)
sub(
    doc,
    "Quote And Login Vertical Position",
    "#loggedOutNav",
    """#loggedOutNav,
#loggedInNav,
.quote-of-day,
#navLoginBtn {
    transform: translateY(2px) !important;
}""",
    [
        "Increase translateY to move the quote and login button down.",
        "Use a negative number like -3px to move them up."
    ],
)
sub(
    doc,
    "Login Button",
    "#navLoginBtn.btn-secondary",
    """#navLoginBtn.btn-secondary {
    min-height: 54px !important;
    padding: 0.85rem 1.35rem !important;
    border-radius: 999px !important;
    border: 1px solid rgba(204, 155, 76, 0.58) !important;
    background: #7d5430 !important;
}""",
    [
        "background changes the button color.",
        "padding controls button width/height.",
        "border-radius controls roundness.",
        "The hover color is in the nearby #navLoginBtn.btn-secondary:hover block."
    ],
)
sub(
    doc,
    "Dormant Ribbon/Menu",
    "Dormant header menu",
    """.bookmark-menu-shell {
    display: none !important;
}

.ribbon-panel.hidden,
.ribbon-panel:not(.is-open) {
    display: none !important;
}""",
    [
        "This keeps the ribbon/dropdown hidden for now.",
        "The JavaScript switch is in script.js: const enableRibbonMenu = false;",
        "Do not delete the markup if you may want this menu later."
    ],
)

doc.add_page_break()
section(doc, "2. Weekly Inspiration", "This section includes the outside box, title, cream prompt area, prompt text, feather button, and video inspiration button.")
sub(
    doc,
    "Outside Box",
    ".compact-inspiration[aria-labelledby=\"weeklyPromptTitle\"]",
    """.compact-inspiration[aria-labelledby="weeklyPromptTitle"] {
    background: var(--nook-page-gradient) fixed !important;
    border-color: var(--nook-panel-border) !important;
}""",
    [
        "This uses the same continuous gradient as the header/footer.",
        "Change border-color for the outline around the box."
    ],
)
sub(
    doc,
    "Weekly Inspiration Title",
    ".compact-inspiration #weeklyPromptTitle",
    """.compact-inspiration .weekly-prompt-heading .section-label-title,
.compact-inspiration #weeklyPromptTitle {
    color: rgba(255, 239, 189, 0.9) !important;
}""",
    [
        "Change color for the Weekly Inspiration label.",
        "Font size may be controlled by the shared .section-label-title or .eyebrow styles."
    ],
)
sub(
    doc,
    "Prompt Text Area / Cream Paper",
    "Weekly Inspiration dream page",
    """.compact-inspiration .prompt-strip {
    background:
        linear-gradient(135deg, rgba(255, 251, 236, 0.98), rgba(232, 213, 171, 0.96)) !important;
    border-color: rgba(179, 141, 82, 0.34) !important;
}""",
    [
        "This is the cream writing-paper area behind the weekly prompt.",
        "To make it flatter, replace the background with one color, like background: #f6ecd2 !important;"
    ],
)
sub(
    doc,
    "Weekly Prompt Text",
    "Readable storybook text",
    """.compact-inspiration #weeklyPromptText {
    font-family: 'Cormorant Garamond', Georgia, serif !important;
    font-weight: 600 !important;
    color: #15100b !important;
    line-height: 1.48 !important;
    font-size: clamp(1.2rem, 1.7vw, 1.42rem) !important;
}""",
    [
        "color changes the text color.",
        "font-weight makes it lighter or bolder; try 500, 600, or 700.",
        "font-size controls the prompt text size."
    ],
)
sub(
    doc,
    "Feather Button Position And Icon Controls",
    "Premium marble feather button",
    """.compact-inspiration {
    --feather-icon-size: 1.9rem;
    --feather-icon-x: -7px;
    --feather-icon-y: -1px;
    --feather-icon-rotation: -10deg;
    --feather-icon-hover-x: 2px;
    --feather-icon-hover-y: -2px;
    --feather-icon-hover-scale: 1.08;
    --feather-icon-hover-rotation: -16deg;
}""",
    [
        "--feather-icon-x moves the feather left/right.",
        "--feather-icon-y moves it up/down.",
        "--feather-icon-size changes icon size.",
        "The hover variables control how it moves on hover."
    ],
)
sub(
    doc,
    "Feather Button Marble Surface",
    ".prompt-write-btn.feather-prompt-btn",
    """.compact-inspiration .prompt-write-btn.feather-prompt-btn,
#copyPromptBtn.prompt-write-btn.feather-prompt-btn {
    border-color: rgba(204, 155, 76, 0.56) !important;
    background:
        radial-gradient(circle at 32% 24%, rgba(255, 229, 158, 0.3), transparent 18%),
        radial-gradient(circle at 63% 70%, rgba(129, 76, 30, 0.66), transparent 32%),
        linear-gradient(145deg, #4b2f1b 0%, #24170f 48%, #080706 100%) !important;
}""",
    [
        "This controls the marble/premium look.",
        "The shine is in the ::after rule and @keyframes featherMarbleShine."
    ],
)
sub(
    doc,
    "Tap For Video Inspiration Button",
    "Video inspiration button",
    """.compact-inspiration .video-toggle-btn,
#toggleVideoBtn.video-toggle-btn {
    min-height: 54px !important;
    padding: 0.85rem 1.45rem !important;
    border-radius: 999px !important;
    background: #7d5430 !important;
}""",
    [
        "The icon was removed in index.html.",
        "Change background for the normal color.",
        "Change the nearby hover block for the hover color."
    ],
)

doc.add_page_break()
section(doc, "3. Writing Desk", "This section includes the desk box, title, controls, writing textarea, paper style, focus mode text, buttons, character counter, pen name, and publish button.")
sub(
    doc,
    "Writing Desk Outside Box",
    "#writingZoneSection.writing-zone",
    """#writingZoneSection.writing-zone {
    background: var(--nook-page-gradient) fixed !important;
    border-color: var(--nook-panel-border) !important;
}""",
    [
        "This is controlled with the shared section gradient.",
        "Search for #writingZoneSection.writing-zone if you want to customize only the writing box."
    ],
)
sub(
    doc,
    "Writing Text Area / Paper",
    "#writingZoneSection #mainStoryInput",
    """#writingZoneSection #mainStoryInput,
body.focus-mode #writingZoneSection #mainStoryInput {
    background:
        linear-gradient(135deg, rgba(255, 251, 236, 0.98), rgba(232, 213, 171, 0.96)) !important;
    color: #15100b !important;
}""",
    [
        "This is the main writing paper.",
        "There are several older textarea rules; edit the final matching block near the bottom if possible."
    ],
)
sub(
    doc,
    "Focus Mode Writing Text",
    "Readable storybook text",
    """body.focus-mode #writingZoneSection #mainStoryInput,
body.focus-mode #writingZoneSection textarea#mainStoryInput {
    font-family: 'Cormorant Garamond', Georgia, serif !important;
    font-weight: 600 !important;
    color: #15100b !important;
    line-height: 1.48 !important;
    font-size: clamp(1.45rem, 2vw, 1.85rem) !important;
}""",
    [
        "This is where you make Focus Mode text more readable.",
        "Increase font-weight for bolder text; increase font-size for larger text."
    ],
)
sub(
    doc,
    "Writing Desk Controls",
    ".writing-desk-controls",
    """.writing-desk-controls,
.focus-ambient-controls {
    /* Search these selectors to change Sound, Light, and Focus buttons. */
}""",
    [
        "Sound and Light menus are controlled by #soundEffectsBtn and #lightingEffectsBtn.",
        "Focus button is #focusModeToggleBtn.",
        "Dropdown menus are #soundEffectsMenu and #lightingEffectsMenu."
    ],
)
sub(
    doc,
    "Writing Style Dropdowns",
    ".typing-control select",
    """#writingZoneSection .typing-control select,
body.focus-mode #writingZoneSection .typing-control select {
    /* Change dropdown color, border, padding, and font here. */
}""",
    [
        "Font dropdown affects the writing area through JavaScript variables.",
        "Size dropdown changes writing scale."
    ],
)
sub(
    doc,
    "Publish Button And Pen Name",
    "#publishBtn",
    """#publishBtn {
    /* Search this ID to change the Publish Story button. */
}

#guestPenName {
    /* Search this ID to change the Pen Name input. */
}""",
    [
        "These are in the Writing Desk bottom controls.",
        "If styling is overridden, add a final block near the bottom of style.css."
    ],
)

section(doc, "4. Reader Favorites", "This section shows top stories or an empty-state message.")
sub(
    doc,
    "Reader Favorites Box",
    ".leaderboard.card",
    """.leaderboard.card {
    background: var(--nook-page-gradient) fixed !important;
    border-color: var(--nook-panel-border) !important;
}""",
    [
        "Change this if you want Reader Favorites to differ from the other boxes.",
        "The heading is #topStoriesTitle in index.html."
    ],
)
sub(
    doc,
    "Top Story Cards",
    ".top-story-card",
    """.top-story-card,
.story-card {
    background: rgba(8, 10, 12, 0.24) !important;
    border-color: rgba(218, 190, 126, 0.18) !important;
}""",
    [
        "This controls individual story cards in Reader Favorites and Latest Submissions.",
        "Change background for card fill; change border-color for outline."
    ],
)

section(doc, "5. Latest Submissions", "This section includes the latest stories, search input, sort dropdown, refresh, load more, and story cards.")
sub(
    doc,
    "Latest Submissions Box",
    ".feed.card",
    """.feed.card {
    background: var(--nook-page-gradient) fixed !important;
    border-color: var(--nook-panel-border) !important;
}""",
    [
        "The heading is #feedTitle in index.html.",
        "This section shares the same card treatment as other main boxes."
    ],
)
sub(
    doc,
    "Search, Sort, Refresh, Load More",
    "#storySearch",
    """#storySearch {
    /* Search input */
}

#feedSort {
    /* Sort dropdown */
}

#refreshFeedBtn,
#loadMoreBtn {
    /* Refresh and Load More buttons */
}""",
    [
        "These controls are inside the Latest Submissions section.",
        "If you want all small buttons to match, search for .btn-secondary and .btn-ghost too."
    ],
)
sub(
    doc,
    "Story Cards And Menus",
    ".story-card",
    """.story-card {
    /* Individual story card */
}

.menu-trigger,
.menu-dropdown {
    /* Three-dot story/comment menu */
}""",
    [
        "The story/comment dropdown menus are different from the dormant header ribbon.",
        "Keep these if you still want edit/delete/action menus on stories."
    ],
)

doc.add_page_break()
section(doc, "6. Footer", "The footer includes the footer logo, footer note, footer links, owl feedback button, and legal text.")
sub(
    doc,
    "Footer Box",
    ".site-footer",
    """.site-footer,
footer {
    background: var(--nook-page-gradient) fixed !important;
    border-color: var(--nook-panel-border) !important;
}""",
    [
        "This makes the footer part of the continuous gradient.",
        "To reduce footer height, look for footer padding and .footer-content spacing."
    ],
)
sub(
    doc,
    "Footer Logo And Brand",
    ".footer-brand",
    """.footer-brand .brand-logo-img {
    /* Footer logo image size */
}

.footer-brand span {
    /* The Story Nook text beside footer logo */
}""",
    [
        "The footer logo uses story-nook-main-logo-new-square.png.",
        "The footer brand text is separate from logo 2."
    ],
)
sub(
    doc,
    "Footer Links",
    ".footer-nav",
    """.footer-nav button,
.footer-links button {
    /* Footer link buttons */
}""",
    [
        "This covers About Us, Guidelines, Coming Soon, Back to Top, etc.",
        "If you remove a footer link visually, remove or hide the button in index.html."
    ],
)
sub(
    doc,
    "Send Us An Owl Button",
    "#footerFeedbackBtn",
    """#footerFeedbackBtn,
.footer-owl-link {
    /* Owl feedback button */
}""",
    [
        "Change button color, border, width, and spacing here.",
        "The owl icon is .footer-owl-icon and the mail icon is .footer-mail-icon."
    ],
)

section(doc, "7. Spacing Between Sections", "Use these rules to adjust vertical spacing without accidentally changing every box.")
sub(
    doc,
    "Page Width And General Box Shape",
    "Box shape and width polish",
    """main {
    max-width: 1120px !important;
    width: min(100% - 36px, 1120px) !important;
}

main > .card,
main > section.card {
    border-radius: 14px !important;
}""",
    [
        "max-width controls how wide the whole content column can get.",
        "width: min(100% - 36px, 1120px) controls side margins on smaller screens.",
        "border-radius controls the sharp/rounded box shape."
    ],
)
sub(
    doc,
    "Space Between Header And Weekly Inspiration Only",
    "main {",
    """main {
    padding-top: 2rem; /* increase this for more space below header */
}""",
    [
        "Use padding-top on main when you only want to change the gap after the header.",
        "This will not change spacing between Weekly Inspiration, Writing Desk, Reader Favorites, and Latest Submissions.",
        "If a later main rule overrides it, add this near the bottom of style.css."
    ],
)
sub(
    doc,
    "Space Between All Main Boxes",
    "main > .card",
    """main > .card,
main > section.card {
    margin-bottom: 1.25rem;
}""",
    [
        "Use this when you want all boxes to have more or less vertical spacing.",
        "Increase margin-bottom for more space between boxes.",
        "Decrease it for a tighter layout."
    ],
)
sub(
    doc,
    "Space Around One Specific Box",
    "#writingZoneSection",
    """#writingZoneSection {
    margin-top: 1.25rem !important;
    margin-bottom: 1.25rem !important;
}""",
    [
        "Use this approach when only one section needs moving.",
        "Weekly Inspiration: .compact-inspiration[aria-labelledby=\"weeklyPromptTitle\"]",
        "Writing Desk: #writingZoneSection",
        "Reader Favorites: .leaderboard.card",
        "Latest Submissions: .feed.card"
    ],
)
add_tip(doc, "Best practice: for one-off spacing changes, target the one exact section. For a global rhythm change, edit main > .card / main > section.card. For only the first gap under the header, edit main padding-top.")

section(doc, "8. Quick File Map", "Where the visible elements are defined.")
sub(
    doc,
    "index.html Elements",
    None,
    """Header: lines 17-29
Weekly Inspiration: lines 76-90
Writing Desk: starts line 97
Reader Favorites: line 265
Latest Submissions: line 272
Footer: lines 295-321""",
    [
        "Use index.html when you want to remove, rename, or reorder visible elements.",
        "Use style.css when you only want to change appearance."
    ],
)
sub(
    doc,
    "Current Important Image Filenames",
    None,
    """Main logo: story-nook-main-logo-new-square.png
Logo 2: story-nook-logo2.png
Background image: story-nook-default-background.jpg""",
    [
        "GitHub Pages is case-sensitive. The filename in index.html must match the uploaded filename exactly.",
        "Avoid capital/lowercase mismatches like Story-nook-logo2.png vs story-nook-logo2.png."
    ],
)

doc.save(OUT)
print(OUT)
